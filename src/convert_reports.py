import re
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """ 設定單格背景顏色 """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC", sz="4", val="single"):
    """ 設定單格邊框 """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)

def add_inline_formatting(paragraph, text, is_code=False, is_quote=False):
    """
    解析粗體 **bold** 和 行內代碼 `code`
    並加入到段落中
    """
    # 正則表達式匹配 **bold** 和 `code`
    pattern = re.compile(r'(\*\*.*?\*\*|`.*?`)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            if is_quote:
                run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(199, 37, 78) # 經典粉紅代碼色
        else:
            if part:
                run = paragraph.add_run(part)
                if is_quote:
                    run.italic = True
                if is_code:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)

def add_code_block(doc, code_lines):
    """ 加入精美的程式碼區塊 """
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # 建立一個 1x1 的表格來包裝程式碼，實現背景色與邊框
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell = table.cell(0, 0)
    set_cell_background(cell, "F5F5F5") # 淺灰色背景
    set_cell_borders(cell, color="CCCCCC", sz="4")
    
    # 寫入程式碼
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.space_before = Pt(4)
    cell_p.paragraph_format.space_after = Pt(4)
    cell_p.paragraph_format.line_spacing = 1.15
    
    code_text = "\n".join(code_lines)
    run = cell_p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(51, 51, 51)

def add_callout(doc, quote_lines, alert_type="NOTE"):
    """ 加入 GitHub 風格的警示/引用區塊 """
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # 根據不同警示類型設定左邊框顏色與背景色
    # NOTE: 藍色, TIP: 綠色, IMPORTANT: 紫色, WARNING: 橘色
    colors = {
        "NOTE": {"border": "2196F3", "bg": "F3F9FE"},
        "TIP": {"border": "4CAF50", "bg": "F4FAF4"},
        "IMPORTANT": {"border": "9C27B0", "bg": "FAF4FC"},
        "WARNING": {"border": "FF9800", "bg": "FFF9F2"},
    }
    
    cfg = colors.get(alert_type, colors["NOTE"])
    
    # 設定左邊框粗，其餘無邊框
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{cfg['border']}"/>
            <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>
            <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>
        </w:tcBorders>
    ''')
    tcPr.append(tcBorders)
    set_cell_background(cell, cfg["bg"])
    
    cell_p = cell.paragraphs[0]
    cell_p.paragraph_format.left_indent = Inches(0.1)
    cell_p.paragraph_format.space_before = Pt(4)
    cell_p.paragraph_format.space_after = Pt(4)
    
    # 加入類型前綴
    prefix_run = cell_p.add_run(f"[{alert_type}] ")
    prefix_run.bold = True
    if alert_type == "NOTE":
        prefix_run.font.color.rgb = RGBColor(33, 150, 243)
    elif alert_type == "TIP":
        prefix_run.font.color.rgb = RGBColor(76, 175, 80)
    elif alert_type == "IMPORTANT":
        prefix_run.font.color.rgb = RGBColor(156, 39, 176)
    elif alert_type == "WARNING":
        prefix_run.font.color.rgb = RGBColor(255, 152, 0)
        
    quote_text = "\n".join(quote_lines)
    add_inline_formatting(cell_p, quote_text, is_quote=True)

def render_table(doc, raw_table_rows):
    """ 解析並繪製 Markdown 表格 """
    if not raw_table_rows:
        return
        
    # 分離出每一格
    parsed_rows = []
    for r in raw_table_rows:
        # 去除頭尾的 | 並分割
        cols = [col.strip() for col in r.strip().split('|')[1:-1]]
        parsed_rows.append(cols)
        
    if not parsed_rows:
        return
        
    # 取得最大列數
    col_count = max(len(row) for row in parsed_rows)
    row_count = len(parsed_rows)
    
    # 建立 Word 表格
    table = doc.add_table(rows=row_count, cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for row_idx, row_data in enumerate(parsed_rows):
        row = table.rows[row_idx]
        
        # 設定行高與緊湊排版
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement('w:trHeight')
        trHeight.set(qn('w:val'), '360')
        trHeight.set(qn('w:hRule'), 'atLeast')
        trPr.append(trHeight)
        
        for col_idx, text in enumerate(row_data):
            if col_idx >= col_count:
                break
            cell = row.cells[col_idx]
            cell.text = "" # 清空預設內容
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            # 第一行是表頭
            if row_idx == 0:
                set_cell_background(cell, "003366") # 深藍色表頭
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # 白色文字
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 隔行著色 (Zebra striping)
                if row_idx % 2 == 0:
                    set_cell_background(cell, "F2F5F8") # 淺灰藍色
                else:
                    set_cell_background(cell, "FFFFFF")
                    
                add_inline_formatting(p, text)
                # 數值靠右，文字靠左
                if text.replace('.', '', 1).isdigit() or "%" in text or "$" in text:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
            set_cell_borders(cell, color="DDDDDD", sz="4")

def markdown_to_docx(md_path, docx_path, title_text):
    """
    將 Markdown 轉換成格式極其精美的 Word (.docx) 檔案
    """
    print(f"正在轉換: {md_path} -> {docx_path}")
    doc = Document()
    
    # 設定頁邊距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # 設定預設字型
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft JhengHei' # 微軟正黑體
    font.size = Pt(11)
    
    # 建立大標題
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(18)
    title_run = title_p.add_run(title_text)
    title_run.font.size = Pt(22)
    title_run.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102) # 皇家深藍色
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    code_lines = []
    
    in_table = False
    table_rows = []
    
    in_blockquote = False
    blockquote_lines = []
    alert_type = "NOTE"
    
    for line in lines:
        stripped = line.strip()
        
        # 1. 處理程式碼區塊
        if stripped.startswith("```"):
            if in_code_block:
                # 程式碼區塊結束，寫入
                add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_lines.append(line.rstrip('\n'))
            continue
            
        # 2. 處理 Markdown 表格
        if stripped.startswith("|"):
            # 排除表頭的分隔線 |---|---|
            if re.match(r'^\|[\s:-|]+$', stripped):
                continue
            in_table = True
            table_rows.append(line)
            continue
        else:
            if in_table:
                render_table(doc, table_rows)
                table_rows = []
                in_table = False
                
        # 3. 處理 GitHub 警示或引用區塊
        if stripped.startswith(">"):
            in_blockquote = True
            # 去除前綴 >
            content = stripped[1:].strip()
            # 檢查是否是 GitHub 警示框標記
            if content.startswith("[!NOTE]"):
                alert_type = "NOTE"
                content = content.replace("[!NOTE]", "").strip()
            elif content.startswith("[!TIP]"):
                alert_type = "TIP"
                content = content.replace("[!TIP]", "").strip()
            elif content.startswith("[!IMPORTANT]"):
                alert_type = "IMPORTANT"
                content = content.replace("[!IMPORTANT]", "").strip()
            elif content.startswith("[!WARNING]"):
                alert_type = "WARNING"
                content = content.replace("[!WARNING]", "").strip()
                
            if content:
                blockquote_lines.append(content)
            continue
        else:
            if in_blockquote:
                add_callout(doc, blockquote_lines, alert_type)
                blockquote_lines = []
                in_blockquote = False
                alert_type = "NOTE"
                
        # 4. 處理空行
        if not stripped:
            continue
            
        # 5. 處理標題 H1, H2, H3, H4
        if stripped.startswith("#"):
            level = 0
            while stripped.startswith("#"):
                level += 1
                stripped = stripped[1:]
            stripped = stripped.strip()
            
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            run = p.add_run(stripped)
            run.bold = True
            
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
                run.font.size = Pt(16)
                run.font.color.rgb = RGBColor(0, 51, 102)
                # 底線效果
                pBdr = OxmlElement('w:pBdr')
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')
                bottom.set(qn('w:sz'), '12')
                bottom.set(qn('w:space'), '4')
                bottom.set(qn('w:color'), '003366')
                pBdr.append(bottom)
                p._p.get_or_add_pPr().append(pBdr)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                run.font.size = Pt(13)
                run.font.color.rgb = RGBColor(0, 102, 204)
            else:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                run.font.size = Pt(11.5)
                run.font.color.rgb = RGBColor(51, 51, 51)
            continue
            
        # 6. 處理無序列表 (*, -)
        if stripped.startswith("* ") or stripped.startswith("- "):
            bullet_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_inline_formatting(p, bullet_text)
            continue
            
        # 7. 處理有序列表
        if re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^(\d+)\.\s(.*)', stripped)
            num = match.group(1)
            list_text = match.group(2).strip()
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            add_inline_formatting(p, list_text)
            continue
            
        # 8. 處理一般段落
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        # 段落首行縮進 (選用)
        # p.paragraph_format.first_line_indent = Inches(0.25)
        add_inline_formatting(p, stripped)
        
    # 收尾處理遺留的表格或代碼塊
    if in_table:
        render_table(doc, table_rows)
    if in_code_block:
        add_code_block(doc, code_lines)
    if in_blockquote:
        add_callout(doc, blockquote_lines, alert_type)
        
    doc.save(docx_path)
    print(f"成功儲存 DOCX 至: {docx_path}")

if __name__ == "__main__":
    src_dir = r"c:\Users\milo9\Desktop\aws112021136\src"
    
    # 轉換報告 1: AWS大數據架構效能比較與雲端費用優化報告
    md_1 = os.path.join(src_dir, "architecture_and_cost_report.md")
    docx_1 = os.path.join(src_dir, "architecture_and_cost_report.docx")
    markdown_to_docx(md_1, docx_1, "AWS 大數據架構效能比較與雲端費用優化報告")
    
    # 轉換報告 2: AWS Lambda 伺服器無缝自動化與 Cloud9 API 串接架構報告
    md_2 = os.path.join(src_dir, "lambda_automation_and_api_report.md")
    docx_2 = os.path.join(src_dir, "lambda_automation_and_api_report.docx")
    markdown_to_docx(md_2, docx_2, "AWS Lambda 伺服器無缝自動化與 Cloud9 API 串接架構報告")
    print("\n[SUCCESS] 所有報告均已成功匯出為高品質 Word 文件！")
